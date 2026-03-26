from flask import Flask, request, send_file, render_template
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, re, zipfile

app = Flask(__name__)

FONT_NAME   = "Arial"
FONT_SIZE   = Pt(12)

LABEL_TAB   = 630
BODY_TAB    = 630
SUB_TAB     = 1260
MONO_INDENT = 5960

TEMPLATE = {
    '[Content_Types].xml': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/><Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/><Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/></Types>',
    '_rels/.rels': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>',
    'word/_rels/document.xml.rels': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/></Relationships>',
    'word/document.xml': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p/><w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr></w:body></w:document>',
    'word/styles.xml': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial" w:eastAsia="Arial"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr></w:pPrDefault></w:docDefaults><w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial" w:eastAsia="Arial"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:style></w:styles>',
    'word/settings.xml': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:defaultTabStop w:val="720"/><w:compat><w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"/><w:compatSetting w:name="overrideTableStyleFontSizeAndJustification" w:uri="http://schemas.microsoft.com/office/word" w:val="1"/><w:compatSetting w:name="enableOpenTypeFeatures" w:uri="http://schemas.microsoft.com/office/word" w:val="1"/><w:compatSetting w:name="doNotFlipMirrorIndents" w:uri="http://schemas.microsoft.com/office/word" w:val="1"/></w:compat></w:settings>',
    'word/theme/theme1.xml': '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Office Theme"><a:themeElements><a:clrScheme name="Office"><a:dk1><a:sysClr val="windowText" lastClr="000000"/></a:dk1><a:lt1><a:sysClr val="window" lastClr="FFFFFF"/></a:lt1><a:dk2><a:srgbClr val="1F497D"/></a:dk2><a:lt2><a:srgbClr val="EEECE1"/></a:lt2><a:accent1><a:srgbClr val="4F81BD"/></a:accent1><a:accent2><a:srgbClr val="C0504D"/></a:accent2><a:accent3><a:srgbClr val="9BBB59"/></a:accent3><a:accent4><a:srgbClr val="8064A2"/></a:accent4><a:accent5><a:srgbClr val="4BACC6"/></a:accent5><a:accent6><a:srgbClr val="F79646"/></a:accent6><a:hlink><a:srgbClr val="0000FF"/></a:hlink><a:folHlink><a:srgbClr val="800080"/></a:folHlink></a:clrScheme><a:fontScheme name="Office"><a:majorFont><a:latin typeface="Arial"/><a:ea typeface=""/><a:cs typeface=""/></a:majorFont><a:minorFont><a:latin typeface="Arial"/><a:ea typeface=""/><a:cs typeface=""/></a:minorFont></a:fontScheme><a:fmtScheme name="Office"><a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="9525"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln><a:ln w="9525"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln><a:ln w="9525"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle><a:effectStyle><a:effectLst/></a:effectStyle><a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:solidFill><a:schemeClr val="phClr"/></a:solidFill><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme></a:themeElements></a:theme>',
}


def make_template():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as z:
        for name, content in TEMPLATE.items():
            z.writestr(name, content)
    buf.seek(0)
    return buf


def set_spacing(para, line=240, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"),     str(line))
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:before"),   str(before))
    sp.set(qn("w:after"),    str(after))
    pPr.append(sp)


def set_tab_stop(para, pos):
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    t = OxmlElement("w:tab")
    t.set(qn("w:val"), "left")
    t.set(qn("w:pos"), str(pos))
    tabs.append(t)


def set_indent(para, left=0, hanging=0):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"),    str(left))
    ind.set(qn("w:hanging"), str(hanging))


def apply_font(run):
    rPr = run._r.get_or_add_rPr()
    for tag in ["w:rFonts", "w:sz", "w:szCs"]:
        for old in rPr.findall(qn(tag)):
            rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    FONT_NAME)
    rFonts.set(qn("w:hAnsi"),    FONT_NAME)
    rFonts.set(qn("w:cs"),       FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.insert(0, rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(FONT_SIZE.pt * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(FONT_SIZE.pt * 2)))
    rPr.append(sz)
    rPr.append(szCs)


def r(para, text, bold=False, underline=False):
    run = para.add_run(text)
    run.font.bold      = bold
    run.font.underline = underline
    apply_font(run)
    return run


def tab(para):
    run = para.add_run("\t")
    apply_font(run)


def rich(para, text):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r(para, part[2:-2], bold=True)
        else:
            r(para, part)


def build_letter(data):
    classification = data.get("classification", "RESTD")
    mono_lines = [data.get(f"ml{i}", "").strip()
                  for i in range(1, 7) if data.get(f"ml{i}", "").strip()]
    date_str   = data.get("date",  "").strip()
    to1        = data.get("to1",   "").strip()
    to2        = data.get("to2",   "").strip()
    to3        = data.get("to3",   "").strip()
    info       = data.get("info",  "").strip()
    id_        = data.get("id",    "").strip()
    ref        = data.get("ref",   "").strip()
    subj       = data.get("subj",  "").strip()
    body_text  = data.get("body",  "")
    sig_lines  = [data.get(f"s{i}", "").strip()
                  for i in range(1, 5) if data.get(f"s{i}", "").strip()]

    doc = Document(make_template())

    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1))

    # 1. RESTD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, line=240, before=0, after=0)
    r(p, classification, bold=True, underline=True)

    # 2. MONOGRAM
    for line in mono_lines:
        p = doc.add_paragraph()
        set_spacing(p, line=240, before=0, after=0)
        set_indent(p, left=MONO_INDENT, hanging=0)
        r(p, line)

    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p, line=240, before=0, after=0)
        r(p, date_str)

    # Gap before address block
    p = doc.add_paragraph()
    set_spacing(p, line=240, before=160, after=0)

    # 3. ADDRESS BLOCK
    if to1:
        p = doc.add_paragraph()
        set_spacing(p, line=240, before=0, after=0)
        set_tab_stop(p, LABEL_TAB)
        r(p, "To:")
        tab(p)
        r(p, to1)
        if to2:
            p = doc.add_paragraph()
            if to3:
             set_spacing(p, line=240, before=0, after=0)
            else:
             set_spacing(p, line=400, before=0, after=0)
            set_indent(p, left=LABEL_TAB, hanging=0)
            r(p, to2)
        if to3:
            p = doc.add_paragraph()
            set_spacing(p, line=400, before=0, after=0)
            set_indent(p, left=LABEL_TAB, hanging=0)
            r(p, to3)

    if info:
        p = doc.add_paragraph()
        set_spacing(p, line=400, before=0, after=0)
        set_tab_stop(p, LABEL_TAB)
        r(p, "Info:")
        tab(p)
        r(p, info)

    if id_:
        p = doc.add_paragraph()
        set_spacing(p, line=360, before=0, after=0)
        set_tab_stop(p, LABEL_TAB)
        r(p, "ID:")
        tab(p)
        r(p, id_)

    # 4. SUBJ
    p = doc.add_paragraph()
    set_spacing(p, line=360, before=0, after=0)
    set_tab_stop(p, LABEL_TAB)
    r(p, "Subj:")
    tab(p)
    r(p, subj, bold=True, underline=True)

    # 5. REF
    if ref:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, line=360, before=0, after=0)
        set_indent(p, hanging=0)
        r(p, ref)

    # 6. BODY
    for raw in body_text.split("\n"):
        t = raw.strip()
        if not t:
            p = doc.add_paragraph()
            set_spacing(p, line=360, before=0, after=0)
            continue

        sub_m  = re.match(r"^([a-z])\.\s+(.*)", t)
        main_m = re.match(r"^(\d+)\.\s+(.*)", t)

        if sub_m:
            letter, rest = sub_m.group(1), sub_m.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p, line=360, before=0, after=0)
            set_tab_stop(p, SUB_TAB)
            set_indent(p, left=SUB_TAB, hanging=SUB_TAB - LABEL_TAB)
            r(p, letter + ".")
            tab(p)
            rich(p, rest)

        elif main_m:
            num, rest = main_m.group(1), main_m.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p, line=360, before=0, after=0)
            set_tab_stop(p, BODY_TAB)
            set_indent(p, left=0, hanging=0)
            r(p, num + ".")
            tab(p)
            rich(p, rest)

        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p, line=360, before=0, after=0)
            rich(p, t)

    # 7. SIGNATORY
    p = doc.add_paragraph()
    set_spacing(p, line=240, before=1260, after=0)

    for line in sig_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p, line=240, before=0, after=0)
        r(p, line)

    # 8. RESTD at bottom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, line=240, before=240, after=0)
    r(p, classification, bold=True, underline=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    data     = request.form.to_dict()
    buf      = build_letter(data)
    subj     = data.get("subj", "letter")
    safe     = re.sub(r"[^a-zA-Z0-9_\-]", "_", subj)[:40]
    filename = f"letter_{safe}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    print("\n  ✓  Server running  →  open  http://localhost:5000  in your browser\n")
    # app.run(debug=True, port=5000)
    
    import os
app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))